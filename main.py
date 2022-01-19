#Resume Parser
import re
import string
import PyPDF2
import docx
# Importing NLTK for stopword removal and tokenizing
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.util import ngrams
import json
import pdfplumber

#Webscrapping part
from bs4 import BeautifulSoup
import requests
import pprint
import datetime
import sqlite3
from sqlite3 import Error
import geeks_interview

#Gui part
import sys
import os, os.path
import csv
import PyQt5.QtWidgets as qtw
from PyQt5.QtGui import QIcon
from PyQt5.QtWidgets import QMainWindow, QApplication, QWidget, QPushButton, QAction, QLineEdit


class App(QMainWindow):

    def __init__(self):
        super().__init__()
        self.title = 'AceTheRace'
        self.left = 10
        self.top = 10
        self.width = 1000
        self.height = 800
        self.initUI()

    def initUI(self):
        self.setWindowTitle(self.title)
        self.setGeometry(self.left, self.top, self.width, self.height)

        # All window components comes here
        self.menu()
        self.mainWidget()

        self.show()

    def menu(self):
        mainMenu = self.menuBar()
        fileMenu = mainMenu.addMenu('File')
        editMenu = mainMenu.addMenu('Edit')
        viewMenu = mainMenu.addMenu('View')
        searchMenu = mainMenu.addMenu('Search')
        toolsMenu = mainMenu.addMenu('Tools')
        helpMenu = mainMenu.addMenu('Help')

        exitSubMenu = QAction(QIcon('exit24.png'), 'Exit', self)
        exitSubMenu.setShortcut('Ctrl+Q')
        exitSubMenu.setStatusTip('Exit application')
        exitSubMenu.triggered.connect(self.close)
        fileMenu.addAction(exitSubMenu)

        newFileMenu = QAction('New', self)
        newFileMenu.setShortcut('Ctrl+N')
        newFileMenu.triggered.connect(self.__init__)
        fileMenu.addAction(newFileMenu)

    def mainWidget(self):
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        vertical = qtw.QVBoxLayout()
        self.centralWidget().setLayout(vertical)
        self.centralWidget().setStyleSheet("background:#000678;")

        # all the widgets comes here

        self.csvLabel = qtw.QLabel("RESUME PARSER", self)
        self.csvLabel.setMinimumSize(300,60)
        self.csvLabel.move(150,75)
        self.csvLabel.setStyleSheet("color:#146356;font-size:24px;")

        self.textbox = QLineEdit(self)
        self.textbox.setPlaceholderText("Enter the the file path")
        self.textbox.setStyleSheet("text-align:center;padding:0px 0px 0px 65px; color:black;")
        self.textbox.move(100, 160)
        self.textbox.setFixedWidth(300)
        self.textbox.setFixedHeight(60)
        #
        self.button1 = QPushButton(self)   # Note since there is self inside the QPushButton(self) it comes above all the widgets
        self.button1.setText("Get Content") # Actually these are not on the innerframe
        self.button1.move(100, 260)
        self.button1.setStyleSheet("padding:10px 20px;text-align:center;background:#f34567;")
        self.button1.setFixedWidth(300)
        self.button1.setFixedHeight(60)
        self.button1.clicked.connect(self.resumeParser)

        # CSV widgets comes here
        self.csvLabel = qtw.QLabel("CSV GENERATOR", self)
        self.csvLabel.setMinimumSize(300,60)
        self.csvLabel.move(650,120)
        self.csvLabel.setStyleSheet("color:#146356;font-size:24px;")

        self.csvButton1 = qtw.QPushButton(self)
        self.csvButton1.setText("GET CSV")
        self.csvButton1.setStyleSheet("background:#f34567;")
        self.csvButton1.setMinimumSize(300,60)
        self.csvButton1.move(600,200)
        self.csvButton1.clicked.connect(self.showdialog)

        #dlg box widgets

        self.dlgTextEdit1 = qtw.QTextEdit()
        self.dlgTextEdit1.setPlaceholderText("Enter the content in normal form, we convert it into csv")
        self.dlgTextEdit1.setMinimumSize(600, 400)
        # dlgTextEdit1.setText("This is is the default text!")

        self.dlgFileName = qtw.QLineEdit()
        self.dlgFileName.setPlaceholderText("Enter the file name to save")
        self.dlgFileName.setMinimumSize(300,60)
        self.dlgFileName.setStyleSheet("background:#e7d1a7; text-align:center")

        self.dlgFilePath = qtw.QLineEdit()
        self.dlgFilePath.setPlaceholderText("Enter the file path to save")
        self.dlgFilePath.setMinimumSize(300,60)
        self.dlgFilePath.setStyleSheet("background:#e7d1a7; text-align:center")


        self.dlgButton = QPushButton()
        self.dlgButton.setMinimumSize(200,60)
        self.dlgButton.setText("Generate CSV")
        self.dlgButton.setStyleSheet("padding:10px 20px;text-align:center;background:#f34567;")
        self.dlgButton.clicked.connect(self.csvGenerator)

        #webscrapping and selenium widgets

        self.webTitle1Label = qtw.QLabel(self)
        self.webTitle1Label.setText("PREPARE FOR YOUR DREAM")
        self.webTitle1Label.setStyleSheet("color:#1A374D; font-size:24px;padding:0px 0px 0px 10px")
        self.webTitle1Label.setMinimumSize(400, 50)
        self.webTitle1Label.move(70, 380)

        # self.webTitle2Label = qtw.QLabel(self)
        # # self.webTitle2Label.setText("")
        # self.webTitle2Label.setStyleSheet("color:black; font-size:24px;")
        # self.webTitle2Label.setMinimumSize(400,50)
        # self.webTitle2Label.move(30,490)

        self.usrNameBox = qtw.QLineEdit()
        self.usrNameBox.setPlaceholderText("Enter the username of LeetCode")
        self.usrNameBox.setMinimumSize(400, 60)

        self.passwdBox = qtw.QLineEdit()
        self.passwdBox.setPlaceholderText("Enter the password of LeetCode")
        self.passwdBox.setMinimumSize(400, 60)

        self.amazonBtn = qtw.QPushButton()
        self.amazonBtn.setText("Amazon")
        self.amazonBtn.setStyleSheet("background:#CDDEFF;font-size:20px;")
        self.amazonBtn.setFixedWidth(200)
        # self.amazonBtn.setMinimumSize(200, 40)
        # self.amazonBtn.move(200, 400)
        self.amazonBtn.clicked.connect(self.amazonTrigger)

        self.flipkartBtn = qtw.QPushButton()
        self.flipkartBtn.setText("Flipkart")
        self.flipkartBtn.setStyleSheet("background:#676FA3;font-size:20px;")
        self.flipkartBtn.setFixedWidth(200)
        # self.flipkartBtn.setMinimumSize(200, 40)
        # self.flipkartBtn.move(200, 500)
        self.flipkartBtn.clicked.connect(self.flipkartTrigger)

        self.adobeBtn = qtw.QPushButton()
        self.adobeBtn.setText("Adobe")
        self.adobeBtn.setStyleSheet("background:#676FA3;font-size:20px;")
        self.adobeBtn.setFixedWidth(200)
        # self.adobeBtn.setMinimumSize(200, 40)
        # self.adobeBtn.move(200, 600)
        self.adobeBtn.clicked.connect(self.adobeTriggger)

        self.googleBtn = qtw.QPushButton()
        self.googleBtn.setText("Google")
        self.googleBtn.setStyleSheet("background:#CDDEFF;font-size:20px;")
        self.googleBtn.setFixedWidth(200)
        # self.googleBtn.setMinimumSize(200, 40)
        # self.googleBtn.move(200, 700)
        self.googleBtn.clicked.connect(self.googleTrigger)

        # *************** All the frames are here *************************

        #Main frames are here

        #resume frames
        resumeFrame, webFrame = qtw.QFrame(), qtw.QFrame()

        resumeFrame.setStyleSheet("background:#124566;text-align:center;")
        resumeLayout = qtw.QGridLayout()
        resumeFrame.setLayout(resumeLayout)

        # Inner Frames are here
        resumeLeftFrame, resumeRightFrame = qtw.QFrame(), qtw.QFrame()
        resumeInnerLayout_L = qtw.QGridLayout()
        resumeInnerLayout_R = qtw.QGridLayout()

        resumeLeftFrame.setLayout(resumeInnerLayout_L)
        resumeLeftFrame.setGeometry(0,0,400,380)
        resumeLeftFrame.setStyleSheet("background:#63d8d6")
        # resumeLeftFrame.setFixedWidth(500)
        # resumeLeftFrame.setFixedHeight(500)
        # resumeInnerLayout_L.addWidget(self.textbox, 1, 1)
        # resumeInnerLayout_L.addWidget(self.button1, 2, 1)

        resumeRightFrame.setLayout(resumeInnerLayout_R)
        resumeRightFrame.setGeometry(0, 0, 400, 380)
        resumeRightFrame.setStyleSheet("background:#43d875")
        # resumeLeftFrame.setMinimumSize(300, 300)
        # resumeInnerLayout_R.addWidget(self.textbox, 1, 1)
        # resumeInnerLayout_R.addWidget(self.button1, 2, 1)

        # Add to resumeLayout
        resumeLayout.addWidget(resumeLeftFrame, 1, 1)
        resumeLayout.addWidget(resumeRightFrame, 1, 2)

        #Web frame
        webFrame.setStyleSheet("background:#f14365;padding:20px;")
        webLayout = qtw.QGridLayout()
        webFrame.setLayout(webLayout)

        webTopFrame, webBottomFrame = qtw.QFrame(), qtw.QFrame()
        webInnerLayout_T = qtw.QGridLayout()
        webInnerLayout_B = qtw.QGridLayout()

        webTopFrame.setLayout(webInnerLayout_T)
        webTopFrame.setStyleSheet("background:#f34567")

        webBottomFrame.setLayout(webInnerLayout_B)
        webBottomFrame.setStyleSheet("background:#f56255")

        #Add wigets to the bottom and the top frames of web frame

        webInnerLayout_T.addWidget(self.usrNameBox, 0,0)
        webInnerLayout_T.addWidget(self.passwdBox,1,0)

        webInnerLayout_B.addWidget(self.amazonBtn,0,1)
        webInnerLayout_B.addWidget(self.adobeBtn,0,2)
        webInnerLayout_B.addWidget(self.flipkartBtn,0,3)
        webInnerLayout_B.addWidget(self.googleBtn,0,4)

        # Add to webLayout original
        webLayout.addWidget(webTopFrame,1,1)
        webLayout.addWidget(webBottomFrame,2,1)




        # Add to webLayout
        # webLayout.addWidget(self.webTitle1Label, 1, 1)
        # webLayout.addWidget(self.usrNameBox,1,2)
        # webLayout.addWidget(self.passwdBox,2,2)
        # webLayout.addWidget(self.amazonBtn, 3,1)
        # webLayout.addWidget(self.flipkartBtn, 3,2)
        # webLayout.addWidget(self.adobeBtn, 3,3)
        # webLayout.addWidget(self.googleBtn, 3,4)
        # webLayout.addWidget(self.googleBtn, 4,1)


        # Add to main frame that is central widget
        vertical.addWidget(resumeFrame)
        vertical.addWidget(webFrame)

    def showdialog(self):
        layout = qtw.QGridLayout()
        self.dlg = qtw.QDialog(self)
        self.dlg.setWindowTitle("CSV TEXT GENERATOR!")
        self.dlg.setLayout(layout)
        self.dlg.setMinimumSize(800, 600)
        layout.addWidget(self.dlgTextEdit1, 0, 0)
        layout.addWidget(self.dlgFileName, 1,0)
        layout.addWidget(self.dlgFilePath, 2,0)
        layout.addWidget(self.dlgButton, 3, 0)
        self.dlg.exec()

    def csvGenerator(self):
        textFromTedit = self.dlgTextEdit1.toPlainText()
        textList = textFromTedit.split(" ")
        # print(textList)
        newCsvLine = ', '.join(textList)
        self.dlgTextEdit1.setPlainText(newCsvLine)
        filename = self.dlgFileName.text()
        filepath = self.dlgFilePath.text()
        fileFinalPath = filepath+"\\"+filename

        with open(os.path.normpath(fileFinalPath), "w+") as csvf:
            writer = csv.writer(csvf)
            writer.writerow(textList)

    def leetCodeData(self):
        user_name = self.usrNameBox.text()
        password = self.passwdBox.text()
        return user_name, password


    def amazonTrigger(self):
        print("amazon\n")
        company_links = self.webScrapping(0)
        user_name, password = self.leetCodeData()
        geeks_interview.open_geeks(company_links[0], user_name, password)

    def adobeTriggger(self):
        print("adobe\n")
        company_links = self.webScrapping(1)
        user_name, password = self.leetCodeData()
        geeks_interview.open_geeks(company_links[0], user_name, password)


    def flipkartTrigger(self):
        print("flipkart\n")
        company_links = self.webScrapping(2)
        user_name, password = self.leetCodeData()
        geeks_interview.open_geeks(company_links[0], user_name, password)


    def googleTrigger(self):
        print("google\n")
        company_links = self.webScrapping(3)
        user_name, password = self.leetCodeData()
        geeks_interview.open_geeks(company_links[0], user_name, password)




    def resumeParser(self):
        filename = self.textbox.text()

        # Extract text from PDF
        def getPDFContent(path):
            content = ""
            # Load PDF into pyPDF
            pdf = PyPDF2.PdfFileReader(open(path, "rb"))
            # Iterate pages
            for i in range(0, pdf.getNumPages()):
                # Extract text from page and add to content
                pageObj = pdf.getPage(i)
                content += pageObj.extractText() + "\n"
            # Collapse whitespace
            content = " ".join(content.replace(u"\xa0", " ").strip().split())
            return content

        def get_pdf_content_using_pdfplumber(path):
            pdf = pdfplumber.open(path)
            text = ""
            for i in range(0, len(pdf.pages)):
                page = pdf.pages[i]
                text += page.extract_text() + "\n"
            return text

        # Extract text from DOCX
        def getText(filename):
            doc = docx.Document(filename)
            fullText = ""
            for para in doc.paragraphs:
                fullText += para.text
            return fullText

        def save_to_docx(s):
            new_doc = docx.Document()
            new_doc.add_paragraph(s)
            new_doc.save('project_output_doc.docx')

        def json_convert(s):
            file = open('python_json.txt', "wt")
            file.write(s)
            file.close()

        # To store extracted resumes
        resume = ""
        # Select a path to the file - code needs os.path #to be addded
        # filename = input("Enter file name / path : ")

        # Invoking document parsers based on file format

        if filename.endswith(".pdf"):
            resume = get_pdf_content_using_pdfplumber(filename)
        elif filename.endswith(".docx"):
            resume = getText(filename)
        else:
            print("File format is currently not supported")
            exit(0)

        print("processing..... \nplease wait....")

        # print(type(resume))
        # Importing NLTK for stopword removal and tokenizing
        from nltk.tokenize import word_tokenize
        from nltk.corpus import stopwords

        # Tokenizing/ Filtering the resume off stopwords and punctuations
        print("tokenizing the given file ......")
        tokens = word_tokenize(resume)
        punctuations = ['(', ')', ';', ':', '[', ']', ',']
        stop_words = stopwords.words('english')
        # storing the cleaned resume
        filtered = [w for w in tokens if not w in stop_words and not w in string.punctuation]
        print("removing the stop words....\nCleaning the resumes....\nExtracting Text .......")

        # get the name from the resume
        name = str(filtered[0]) + ' ' + str(filtered[1])

        # using regular expressions we extract phone numbers and mail ids

        # get contact info - from resume
        # email
        email = ""
        match_mail = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', resume)
        # handling the cases when mobile number is not given
        if (match_mail != None):
            email = match_mail.group(0)

        # mobile number
        mobile = ""
        match_mobile = re.search(r'[789]\d{9}', resume)
        # handling the cases when mobile number is not given
        if (match_mobile != None):
            mobile = match_mobile.group(0)

        # creating a text file of all keywords in resume and joining them by newline
        parsed_resume = '\n'.join(filtered)

        r = str("Name: " + name + "\n" + "Email : " + email + "\n" + "Mobile :" + mobile + "\n\n" + parsed_resume)

        # saving to docx file
        save_to_docx(r)

        # creating a dictionary to store as json format
        json_string = {"Name": name, "Email": email, "Mobile": mobile, "Keywords": str(parsed_resume)}

        # storing the json string in file

        json_convert(json.dumps(json_string))

    def webScrapping(self, coname):

        companies = ['amazon', 'adobe', 'flipkart', 'google']  # these are companies you can select
        # coname = int(input('enter company name 0:amazon  1:adobe  2:flipkart 3:google '))  # list to be added to the dropdown
        url = "https://www.geeksforgeeks.org/tag/" + companies[coname]  # url to selecyed company
        res = requests.get(url)
        content = BeautifulSoup(res.text, 'html.parser')  # taking content using beautiful soup from the page
        link = content.findAll("div", attrs={"class": "head"})
        li = []
        for l in link:
            try:
                a = l.find('a')['href']
            except:
                continue
            else:
                li.append(a)

        print(pprint.pprint(li))

        comp_name = companies[coname]
        comp_links = li
        comp_dict = {'company_name': [], 'quest_links': []}
        name = comp_dict['company_name'].append(comp_name)
        ques = comp_dict['quest_links'].append(comp_links)


        # connecting to db sqlite
        def sql_connection():
            try:
                con = sqlite3.connect('mydatabase.db')
                return con
            except Error:
                print(Error)

        def sql_table(con):
            cursorObj = con.cursor()

            # cursorObj.execute("CREATE TABLE company3(com_name varchar(20), ques_link str(1000),created_at str(100))")
            cursorObj.execute("INSERT INTO company3 values ( ?  , ? ,?);",(comp_name, str(comp_links), str(datetime.datetime.now())))

            con.commit()

        def selection(con, com_name):
            cursorObj = con.cursor()
            cursorObj.execute("select * from company3 where com_name = " + '"' + com_name + '"' + ";")
            rows = cursorObj.fetchall()
            # print()
            # print("rows variable")
            # print(rows)

            return rows

        # sql_table(con)#uncomment this to run
        # convertion to pdf
        def pdfconvertion(con1):
            sql_table(con1)

            rows = selection(con1, companies[coname])
            doc = docx.Document()
            doc.add_paragraph(
                "company : " + rows[0][0] + "\nInterview Experience Links : " + rows[0][1] + "\nTime: " + rows[0][2])

            doc.save('hello' + comp_name + '.docx')

        con1 = sql_connection()
        # sql_table(con1)
        pdfconvertion(con1)
        return comp_links


if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = App()
    sys.exit(app.exec_())
